"""반입 파일 판별·B군 추출 계층 (S16 — F42, 설계 §4.18) 단위 테스트.

불변 규칙 7의 예외 — "판별 미통과 바이너리의 텍스트 투입 경로 0건"·"cp949 mojibake 차단"·
"파서 예외 구조화"는 반입 정확성에 직결되는 핵심 로직이라 qnet `_magic_ok`·
`_unsupported_message` 테스트 전례와 같은 급의 단위 테스트로 취급한다(stage-16 5절).

고정하는 계약(설계 v1.19 — Opus 검토 적발분 포함):
  ① 매직 바이트 우선 판별(zip 내부 docx/xlsx/hwpx 판별·OLE·PDF·이미지) — 확장자 사칭 방어.
     **판별 결과(group)가 후단 콘텐츠 블록/프롬프트 조립까지 지배**한다 — 확장자 없는
     `%PDF`·PNG가 `path.suffix` 재분기로 텍스트 폴백에 새지 않는다(검토 1).
  ② 인코딩 판별 3단(utf-8 BOM → utf-8 → cp949) — cp949 mojibake 차단(DoD 3)
  ③ C군(zip·hwp·hwpx·doc·xls·판별 불가) → UnsupportedFormatError + 원본 sources/ 저장 +
     포맷별 폴백 문구(§4.18 ⑥). hwpx도 hwp와 동일 문구(검토 3).
  ④ B군(docx·xlsx) 추출 — Markdown 표 직렬화(직렬화 루프가 세는 값으로 상한 판정, 검토 4),
     파서 예외 → DocParseFailedError(parse_failed 재사용) + 원본 sources/ 저장(검토 2).
     암호 걸린 OOXML(OLE 컨테이너)도 parse_failed(검토 7).
  ⑤ 200,000자 상한(A군 디코드·B군 추출 공통, §4.18 ⑤ 확정 문안 그대로 — 검토 6) 및
     xlsx 시트10·행500·열50 상한 → TooLargeError
  ⑥ convert 잡의 unsupported_format·parse_failed는 alternatives=[], fetch 잡은 기존
     기본값 유지(검토 5)
"""
from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook

from services import codex_adapter, convert_service, doc_extract, import_service
from services.fetchers.base import Adapter, FetchedFile

_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


# ---------------------------------------------------------------------------
# 픽스처 — sources/ 실경로를 tmp_path로 갈아끼운다(qnet·trust_gate 테스트 전례).
# ---------------------------------------------------------------------------
@pytest.fixture()
def isolated_sources(tmp_path, monkeypatch):
    sources_dir = tmp_path / "sources"
    monkeypatch.setattr(import_service, "SOURCES_DIR", sources_dir)
    return sources_dir


def _make_docx_bytes(paragraphs=None, table_rows=None) -> bytes:
    doc = DocxDocument()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes(sheets: dict) -> bytes:
    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _plain_zip_bytes() -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("readme.txt", "hello")
    return buf.getvalue()


def _hwpx_bytes(*, via: str = "mimetype") -> bytes:
    """hwpx는 zip이라 내부 마커로만 구분한다(§4.18 ② — 검토 3). `via`로 두 식별 경로를
    각각 재현: 'mimetype' 엔트리 또는 'Contents/' 최상위 구조."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        if via == "mimetype":
            zf.writestr("mimetype", "application/hwp+zip")
            zf.writestr("version.xml", "<hwpx/>")
        else:
            zf.writestr("Contents/content.hpf", "<hwpx/>")
    return buf.getvalue()


# ---------------------------------------------------------------------------
# doc_extract — docx
# ---------------------------------------------------------------------------
def test_extract_docx_text_preserves_order_and_renders_table_as_markdown():
    data = _make_docx_bytes(
        paragraphs=["첫 문단입니다.", "둘째 문단입니다."],
        table_rows=[["열1", "열2"], ["a", "b"]],
    )
    text = doc_extract.extract_docx_text(data)
    assert "첫 문단입니다." in text
    assert text.index("첫 문단입니다.") < text.index("둘째 문단입니다.")
    assert "| 열1 | 열2 |" in text
    assert "| a | b |" in text


def test_extract_docx_text_rejects_corrupt_bytes():
    with pytest.raises(doc_extract.DocExtractError):
        doc_extract.extract_docx_text(b"this is not a docx file at all")


def test_extract_docx_text_rejects_structurally_damaged_zip():
    """word/document.xml은 있지만 패키지 구조가 깨진 경우(암호·손상·위장 시뮬레이션)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("word/document.xml", b"not real xml <<<")
        zf.writestr("[Content_Types].xml", "<Types/>")
    with pytest.raises(doc_extract.DocExtractError):
        doc_extract.extract_docx_text(buf.getvalue())


def test_docx_has_lost_elements_detects_drawing_marker():
    xml = (
        b'<w:document xmlns:w="ns"><w:body><w:p><w:r><w:drawing/></w:r></w:p></w:body>'
        b"</w:document>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("word/document.xml", xml)
    assert doc_extract.docx_has_lost_elements(buf.getvalue()) is True


def test_docx_has_lost_elements_false_for_plain_text_docx():
    data = _make_docx_bytes(paragraphs=["그림도 수식도 없는 문서입니다."])
    assert doc_extract.docx_has_lost_elements(data) is False


# ---------------------------------------------------------------------------
# doc_extract — xlsx
# ---------------------------------------------------------------------------
def test_extract_xlsx_text_serializes_sheet_as_markdown_table_with_cell_escaping():
    data = _make_xlsx_bytes(
        {"Sheet1": [["헤더1", "헤더2"], ["a|b", "줄바꿈\n포함"]]}
    )
    text = doc_extract.extract_xlsx_text(data)
    assert "## 시트: Sheet1" in text
    assert "헤더1" in text and "헤더2" in text
    assert "a\\|b" in text  # 셀 내 파이프 이스케이프(§4.18 ④)
    assert "줄바꿈 포함" in text  # 셀 내 줄바꿈은 공백 치환


def test_extract_xlsx_text_rejects_corrupt_bytes():
    with pytest.raises(doc_extract.DocExtractError):
        doc_extract.extract_xlsx_text(b"this is not an xlsx file at all")


def test_extract_xlsx_text_rejects_too_many_sheets():
    sheets = {f"S{i}": [["a"]] for i in range(11)}
    data = _make_xlsx_bytes(sheets)
    with pytest.raises(doc_extract.DocTooLargeError) as exc_info:
        doc_extract.extract_xlsx_text(data)
    assert "시트" in exc_info.value.public_message
    assert "10" in exc_info.value.public_message


def test_extract_xlsx_text_rejects_row_overflow_without_partial_truncation():
    rows = [["v"] for _ in range(501)]
    data = _make_xlsx_bytes({"Sheet1": rows})
    with pytest.raises(doc_extract.DocTooLargeError) as exc_info:
        doc_extract.extract_xlsx_text(data)
    message = exc_info.value.public_message
    assert "Sheet1" in message and "501" in message and "500" in message


def test_extract_xlsx_text_rejects_col_overflow():
    row = ["v"] * 51
    data = _make_xlsx_bytes({"Sheet1": [row]})
    with pytest.raises(doc_extract.DocTooLargeError) as exc_info:
        doc_extract.extract_xlsx_text(data)
    assert "열" in exc_info.value.public_message and "50" in exc_info.value.public_message


def test_extract_xlsx_text_does_not_false_positive_on_formatting_only_cells():
    """S16 검토 4 — `ws.max_row`(서식만 있는 빈 셀 포함)가 아니라 **실제 값이 있는 행 수**로
    상한을 판정한다. 실측 재현: 데이터 3행 + 아무 값도 없이 서식만 적용된 600번째 행."""
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["a", "b"])
    ws.append(["c", "d"])
    ws.append(["e", "f"])
    ws.cell(row=600, column=1).font = Font(bold=True)  # 값 없이 서식만 — used range를 넓힌다
    buf = io.BytesIO()
    wb.save(buf)
    data = buf.getvalue()

    assert ws.max_row == 600  # 전제 확인 — openpyxl이 실제로 이렇게 리포트한다

    text = doc_extract.extract_xlsx_text(data)  # 예외 없이 통과해야 한다(오탐 수정)
    assert "## 시트: Sheet1" in text
    assert "| a | b |" in text


def test_enforce_max_chars_boundary():
    doc_extract.enforce_max_chars("x" * doc_extract.MAX_TEXT_CHARS)  # 경계값은 통과
    with pytest.raises(doc_extract.DocTooLargeError):
        doc_extract.enforce_max_chars("x" * (doc_extract.MAX_TEXT_CHARS + 1))


def test_enforce_max_chars_uses_confirmed_wording():
    """S16 검토 6 — §4.18 ⑤ 확정 문안을 그대로 쓴다("추출된 텍스트가...", "원본 텍스트가"
    아님)."""
    with pytest.raises(doc_extract.DocTooLargeError) as exc_info:
        doc_extract.enforce_max_chars("x" * (doc_extract.MAX_TEXT_CHARS + 1))
    assert exc_info.value.public_message.startswith("추출된 텍스트가 너무 깁니다")


# ---------------------------------------------------------------------------
# convert_service._detect_import_format — 매직 바이트 우선 판별
# ---------------------------------------------------------------------------
def test_detect_pdf_by_magic_bytes_ignores_extension():
    detected = convert_service._detect_import_format("실제내용.txt", b"%PDF-1.4 fake pdf body")
    assert detected.group == "pdf"


def test_detect_image_png_by_magic_bytes():
    detected = convert_service._detect_import_format("photo.jpg", b"\x89PNG\r\n\x1a\n" + b"rest")
    assert detected.group == "image"
    assert detected.file_type == "png"


def test_detect_text_utf8_bom_decodes_cleanly():
    data = "안녕하세요 — BOM 포함".encode("utf-8-sig")
    detected = convert_service._detect_import_format("note.txt", data)
    assert detected.group == "text"
    assert detected.encoding == "utf-8-sig"
    assert detected.text == "안녕하세요 — BOM 포함"


def test_detect_text_cp949_decodes_without_mojibake():
    """DoD 3 — cp949 텍스트 파일이 깨지지 않고 판별된다(D2-② 3단 폴백)."""
    original = "한국어 기출 문제 반입 검증용 본문입니다."
    data = original.encode("cp949")
    detected = convert_service._detect_import_format("기출.txt", data)
    assert detected.group == "text"
    assert detected.encoding == "cp949"
    assert detected.text == original


@pytest.mark.parametrize(
    "filename,expected_file_type",
    [("a.html", "html"), ("a.htm", "html"), ("a.xml", "xml"), ("a.csv", "csv"), ("a.md", "md")],
)
def test_detect_text_group_file_type_by_extension(filename, expected_file_type):
    detected = convert_service._detect_import_format(filename, "본문".encode("utf-8"))
    assert detected.group == "text"
    assert detected.file_type == expected_file_type


def test_detect_unknown_extension_text_falls_back_to_txt():
    """확장자 미상 텍스트는 txt 취급(§4.18 ②-3) — .docx 이름의 실제 텍스트 파일도 포함."""
    detected = convert_service._detect_import_format("이름없음", "본문".encode("utf-8"))
    assert detected.file_type == "txt"
    fake_docx_name = convert_service._detect_import_format("가짜.docx", "본문".encode("utf-8"))
    assert fake_docx_name.group == "text"
    assert fake_docx_name.file_type == "txt"


def test_detect_docx_zip_marker_extracts_text(isolated_sources):
    data = _make_docx_bytes(paragraphs=["1번 문제. 정답은 무엇인가?"])
    detected = convert_service._detect_import_format("기출.docx", data)
    assert detected.group == "docx"
    assert "1번 문제" in detected.text


def test_detect_xlsx_zip_marker_extracts_text(isolated_sources):
    data = _make_xlsx_bytes({"Sheet1": [["a", "b"]]})
    detected = convert_service._detect_import_format("기출.xlsx", data)
    assert detected.group == "xlsx"
    assert "시트: Sheet1" in detected.text


def test_detect_docx_with_lost_elements_appends_note(isolated_sources, monkeypatch):
    """소실 감지(`docx_has_lost_elements`) 자체는 별도 단위 테스트가 커버 — 여기서는
    `_detect_import_format`이 그 판정을 잡 notes 소표기로 정확히 이어붙이는지만 검증한다."""
    monkeypatch.setattr(doc_extract, "docx_has_lost_elements", lambda data: True)
    data = _make_docx_bytes(paragraphs=["본문"])
    detected = convert_service._detect_import_format("수식포함.docx", data)
    assert detected.notes and "수식" in detected.notes[0]


def test_detect_docx_without_lost_elements_has_no_note(isolated_sources):
    data = _make_docx_bytes(paragraphs=["평범한 문서"])
    detected = convert_service._detect_import_format("평범.docx", data)
    assert detected.notes == []


# ---------------------------------------------------------------------------
# convert_service._detect_import_format — C군 거부(조용한 스킵 금지 + sources/ 저장)
# ---------------------------------------------------------------------------
def test_detect_plain_zip_rejected_and_original_saved(isolated_sources):
    data = _plain_zip_bytes()
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("묶음.zip", data)
    assert "ZIP" in exc_info.value.public_message
    assert exc_info.value.action.startswith("원본은 sources/에 저장했습니다.")
    saved = list(isolated_sources.glob("*묶음.zip"))
    assert len(saved) == 1
    assert saved[0].read_bytes() == data  # 원본 불변 — 저장 바이트는 원 바이트 그대로


def test_detect_hwp_ole_rejected_with_hwp_message(isolated_sources):
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("기출.hwp", data)
    assert "한글" in exc_info.value.public_message
    assert "PDF" in exc_info.value.action


@pytest.mark.parametrize("via", ["mimetype", "contents_dir"])
def test_detect_hwpx_zip_uses_hwp_message_not_generic_zip(isolated_sources, via):
    """S16 검토 3 — hwpx는 zip 계열이지만 "압축을 풀어…" 일반 zip 문구가 아니라 hwp와
    동일한 "한글에서 PDF로…" 문구로 안내돼야 한다(§4.18 ⑥ "hwp·hwpx" 행)."""
    data = _hwpx_bytes(via=via)
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("기출.hwpx", data)
    assert "한글" in exc_info.value.public_message
    assert "PDF" in exc_info.value.action
    assert "압축" not in exc_info.value.public_message  # 일반 zip 문구로 오안내되지 않는다


def test_is_hwpx_helper_detects_both_markers():
    for via in ("mimetype", "contents_dir"):
        data = _hwpx_bytes(via=via)
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            assert convert_service._is_hwpx(zf, set(zf.namelist())) is True
    plain = _plain_zip_bytes()
    with zipfile.ZipFile(io.BytesIO(plain)) as zf:
        assert convert_service._is_hwpx(zf, set(zf.namelist())) is False


def test_detect_encrypted_docx_ole_maps_to_parse_failed_not_unsupported(isolated_sources):
    """S16 검토 7 — 암호화된 OOXML(ECMA-376)은 zip이 아니라 OLE CFB 컨테이너. 확장자로
    docx가 **확정**되므로 parse_failed(암호 해제 안내)여야 한다 — unsupported_format이
    아니다(신규 파서 라이브러리 도입 없음, 판별만)."""
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.DocParseFailedError) as exc_info:
        convert_service._detect_import_format("암호걸림.docx", data)
    assert "암호" in exc_info.value.public_message
    assert exc_info.value.action.startswith("원본은 sources/에 저장했습니다.")
    assert "암호" in exc_info.value.action
    saved = list(isolated_sources.glob("*암호걸림.docx"))
    assert len(saved) == 1  # 검토 2 — parse_failed도 원본을 저장한다


def test_detect_encrypted_xlsx_ole_maps_to_parse_failed(isolated_sources):
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.DocParseFailedError) as exc_info:
        convert_service._detect_import_format("암호걸림.xlsx", data)
    assert "엑셀" in exc_info.value.public_message or "암호" in exc_info.value.public_message
    saved = list(isolated_sources.glob("*암호걸림.xlsx"))
    assert len(saved) == 1


def test_detect_doc_ole_rejected_with_doc_message(isolated_sources):
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("기출.doc", data)
    assert "워드" in exc_info.value.public_message
    assert "docx" in exc_info.value.action


def test_detect_xls_ole_rejected_with_xls_message(isolated_sources):
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("기출.xls", data)
    assert "엑셀" in exc_info.value.public_message
    assert "xlsx" in exc_info.value.action


def test_detect_ole_unknown_extension_uses_generic_message(isolated_sources):
    data = _OLE_MAGIC + b"\x00" * 64
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("확장자없음", data)
    assert "구버전" in exc_info.value.public_message


def test_detect_binary_with_null_bytes_is_undetected(isolated_sources):
    data = b"\x00\x01\x02\x03" * 100
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("mystery.bin", data)
    assert "판별할 수 없습니다" in exc_info.value.public_message


def test_detect_undecodable_binary_without_null_bytes_is_undetected(isolated_sources):
    """널 바이트는 없지만 utf-8·cp949 둘 다 실패하는 바이트 — 조용히 텍스트로 흘러가면 안 된다."""
    data = bytes([0xFF, 0xFE, 0x81, 0x8F]) * 20
    with pytest.raises(convert_service.UnsupportedFormatError) as exc_info:
        convert_service._detect_import_format("mystery.dat", data)
    assert "판별할 수 없습니다" in exc_info.value.public_message


# ---------------------------------------------------------------------------
# convert_service._detect_import_format — 200,000자 상한(too_large, LLM 호출 전 차단)
# ---------------------------------------------------------------------------
def test_detect_text_over_char_limit_raises_too_large_before_llm_call():
    huge = ("가" * doc_extract.MAX_TEXT_CHARS) + "초과분"
    with pytest.raises(convert_service.TooLargeError) as exc_info:
        convert_service._detect_import_format("huge.txt", huge.encode("utf-8"))
    assert "200,000" in exc_info.value.public_message
    # S16 검토 6 — §4.18 ⑤ 확정 문안 그대로("원본 텍스트가"가 아니라 "추출된 텍스트가").
    assert exc_info.value.public_message.startswith("추출된 텍스트가 너무 깁니다")
    assert exc_info.value.action  # 분할 권고 문구가 비어 있지 않다


def test_detect_xlsx_row_overflow_maps_to_too_large_error(isolated_sources):
    rows = [["v"] for _ in range(501)]
    data = _make_xlsx_bytes({"Sheet1": rows})
    with pytest.raises(convert_service.TooLargeError) as exc_info:
        convert_service._detect_import_format("big.xlsx", data)
    assert "500" in exc_info.value.public_message


# ---------------------------------------------------------------------------
# convert_service._detect_import_format — B군 파서 예외 → parse_failed 재사용(신규 kind 없음)
# ---------------------------------------------------------------------------
def test_detect_damaged_docx_maps_to_doc_parse_failed_not_unsupported(isolated_sources):
    """판별이 docx로 **확정된 뒤**의 추출 실패는 parse_failed(§4.18 ④ 경계 규칙) —
    unsupported_format이 아니다(판별 자체는 성공했으므로). S16 검토 2 — 원본도 저장한다."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("word/document.xml", b"not real xml <<<")
        zf.writestr("[Content_Types].xml", "<Types/>")
    data = buf.getvalue()
    with pytest.raises(convert_service.DocParseFailedError) as exc_info:
        convert_service._detect_import_format("손상.docx", data)
    assert exc_info.value.action.startswith("원본은 sources/에 저장했습니다.")
    saved = list(isolated_sources.glob("*손상.docx"))
    assert len(saved) == 1
    assert saved[0].read_bytes() == data


def test_detect_damaged_xlsx_maps_to_doc_parse_failed(isolated_sources):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("xl/workbook.xml", b"not real xml <<<")
    data = buf.getvalue()
    with pytest.raises(convert_service.DocParseFailedError) as exc_info:
        convert_service._detect_import_format("손상.xlsx", data)
    assert exc_info.value.action.startswith("원본은 sources/에 저장했습니다.")
    saved = list(isolated_sources.glob("*손상.xlsx"))
    assert len(saved) == 1


# ---------------------------------------------------------------------------
# error_info 매핑 — kind 재사용·신설 계약(§4.11·§4.18) 확인
# ---------------------------------------------------------------------------
def test_fallback_error_info_maps_too_large_kind():
    exc = convert_service.TooLargeError("추출된 텍스트가 너무 깁니다(약 300,000자 — 상한 200,000자)")
    info = convert_service._fallback_error_info(exc, job_kind="convert")
    assert info["kind"] == "too_large"
    assert info["fallback_available"] is False
    assert info["alternatives"] == []


def test_fallback_error_info_maps_doc_parse_failed_to_parse_failed_kind():
    exc = convert_service.DocParseFailedError(
        "워드(docx) 문서를 열 수 없습니다(손상되었거나 암호가 걸려 있을 수 있습니다).",
        action="원본은 sources/에 저장했습니다. 암호를 해제하거나 PDF로 저장한 뒤 다시 반입하세요.",
    )
    info = convert_service._fallback_error_info(exc, job_kind="convert")
    assert info["kind"] == "parse_failed"  # 신규 kind 없음 — 기존 값 재사용
    assert "암호" in info["action"]
    assert info["action"].startswith("원본은 sources/에 저장했습니다.")


def test_fallback_error_info_convert_job_forces_empty_alternatives_for_unsupported_and_parse_failed(
    isolated_sources,
):
    """S16 검토 5(§4.18 ⑥ v1.19 확정) — convert 잡(파일·URL 반입) 발생분은 실패한 반입
    경로 자신을 대안으로 제시하지 않는다(alternatives=[])."""
    unsupported = convert_service._unsupported_error("묶음.zip", _plain_zip_bytes(), "zip")
    info1 = convert_service._fallback_error_info(unsupported, job_kind="convert")
    assert info1["kind"] == "unsupported_format"
    assert info1["alternatives"] == []

    parse_failed = convert_service.DocParseFailedError("워드(docx) 문서를 열 수 없습니다.")
    info2 = convert_service._fallback_error_info(parse_failed, job_kind="convert")
    assert info2["kind"] == "parse_failed"
    assert info2["alternatives"] == []


def test_fallback_error_info_fetch_job_keeps_existing_default_alternatives(isolated_sources):
    """S16 검토 5 — fetch 잡(qnet) 발생분은 기존 기본값을 유지한다(건드리지 않는다)."""
    unsupported = convert_service._unsupported_error("묶음.zip", _plain_zip_bytes(), "zip")
    info1 = convert_service._fallback_error_info(unsupported, job_kind="fetch")
    assert info1["alternatives"] == ["file_import", "url_import"]

    parse_failed = convert_service.DocParseFailedError("워드(docx) 문서를 열 수 없습니다.")
    info2 = convert_service._fallback_error_info(parse_failed, job_kind="fetch")
    assert info2["alternatives"] == ["url_import"]


# ---------------------------------------------------------------------------
# URL 반입 content-type 화이트리스트 확장(§4.18 ⑦, D2-⑤ 확정 5종)
# ---------------------------------------------------------------------------
def test_url_content_type_whitelist_includes_new_five_formats():
    allowed = convert_service._ALLOWED_CONTENT_TYPES
    assert allowed["application/xml"] == "xml"
    assert allowed["text/xml"] == "xml"
    assert allowed["text/csv"] == "csv"
    assert (
        allowed["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
        == "docx"
    )
    assert (
        allowed["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"] == "xlsx"
    )
    assert "application/octet-stream" not in allowed  # 여전히 미허용(YAGNI)


# ---------------------------------------------------------------------------
# S16 검토 1 — 판별 결과(group)가 콘텐츠 블록/프롬프트 조립까지 지배한다(확장자 재분기 없음).
# 검토 실측 재현: 확장자 없는 %PDF·PNG가 utf-8 replace 텍스트 폴백으로 새던 결함.
# ---------------------------------------------------------------------------
def test_api_content_blocks_uses_group_not_extension_for_extensionless_pdf(tmp_path):
    pdf_path = tmp_path / "representative.download"  # 확장자가 pdf가 아님(실측 재현)
    pdf_path.write_bytes(b"%PDF-1.4 fake pdf body")
    blocks = convert_service._api_content_blocks_for_file(pdf_path, group="pdf", file_type="pdf")
    assert len(blocks) == 1
    assert blocks[0]["type"] == "document"
    assert blocks[0]["source"]["media_type"] == "application/pdf"


def test_api_content_blocks_uses_group_not_extension_for_extensionless_png(tmp_path):
    png_path = tmp_path / "attachment.bin"
    png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"rest")
    blocks = convert_service._api_content_blocks_for_file(png_path, group="image", file_type="png")
    assert len(blocks) == 1
    assert blocks[0]["type"] == "image"
    assert blocks[0]["source"]["media_type"] == "image/png"


def test_api_content_blocks_group_pdf_wins_even_if_suffix_says_txt(tmp_path):
    """검토 재현 케이스 그대로 — `.txt`로 끝나는데 내용은 PDF인 파일."""
    path = tmp_path / "실제내용.txt"
    path.write_bytes(b"%PDF-1.4 fake pdf body")
    blocks = convert_service._api_content_blocks_for_file(path, group="pdf", file_type="pdf")
    assert blocks[0]["type"] == "document"  # 텍스트 블록(utf-8 replace)으로 새지 않는다


def test_build_convert_prompt_api_forwards_group_to_content_blocks(tmp_path):
    pdf_path = tmp_path / "no_extension"
    pdf_path.write_bytes(b"%PDF-1.4 fake pdf body")
    convert_md = "지시문"
    _prompt_text, blocks = convert_service._build_convert_prompt_api(
        convert_md, pdf_path, group="pdf", file_type="pdf"
    )
    assert blocks[0]["type"] == "document"


def test_codex_build_text_for_prompt_group_override_ignores_extension(tmp_path, monkeypatch):
    """codex 경로 — `group='pdf'`를 명시하면 확장자(.download)와 무관하게 pypdf 추출
    경로를 탄다(검토 1). 기존 확장자 기반 호출(하위 호환)은 별도 테스트가 이미 고정한다
    (`test_llm_engine_registry.py`)."""
    pdf_path = tmp_path / "representative.download"
    pdf_path.write_bytes(b"%PDF-1.4 fake pdf body")
    monkeypatch.setattr(
        codex_adapter.source_match, "extract_pdf_text", lambda data: "추출된 pdf 텍스트"
    )
    text = codex_adapter.build_text_for_prompt(pdf_path, group="pdf")
    assert text == "추출된 pdf 텍스트"


def test_codex_build_text_for_prompt_group_image_rejects_without_extension(tmp_path):
    from exceptions import ValidationAppError

    png_path = tmp_path / "attachment.bin"
    png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"rest")
    with pytest.raises(ValidationAppError):
        codex_adapter.build_text_for_prompt(png_path, group="image")


def test_codex_build_text_for_prompt_group_text_reads_raw_even_with_pdf_like_extension(tmp_path):
    """group이 'text'/'docx'/'xlsx'면 codex 경로에서 이 함수 자체가 호출되지 않지만(embedded
    경로 우회), 방어적으로 group='text'를 넘기면 pdf/image로 오분류되지 않아야 한다."""
    path = tmp_path / "note.txt"
    path.write_text("안녕하세요", encoding="utf-8")
    text = codex_adapter.build_text_for_prompt(path, group="text")
    assert text == "안녕하세요"


# ---------------------------------------------------------------------------
# S16 검토 8 — LLM 사용량 사전 안내(ETA 표본)를 추출 텍스트 길이 기준으로 정확화
# ---------------------------------------------------------------------------
def test_detection_input_size_uses_text_length_for_text_group():
    detected = convert_service.ImportDetection(group="text", file_type="txt", text="가" * 100)
    assert convert_service._detection_input_size(detected, fallback_bytes=999999) == len(
        ("가" * 100).encode("utf-8")
    )


def test_detection_input_size_falls_back_to_bytes_for_pdf_image():
    detected = convert_service.ImportDetection(group="pdf", file_type="pdf")
    assert convert_service._detection_input_size(detected, fallback_bytes=12345) == 12345


# ---------------------------------------------------------------------------
# S16 검토 9 — cp949 CLI 재인코딩 tmp는 CLI가 실제로 선택됐을 때만 지연 생성한다.
# ---------------------------------------------------------------------------
def test_ensure_cli_recoded_tmp_creates_utf8_file_and_caches(tmp_path, monkeypatch):
    monkeypatch.setattr(convert_service, "CONVERT_TMP_DIR", tmp_path)
    job = {"_extra_tmp_paths": [], "_cli_recoded_path": None}
    text = "한국어 재인코딩 검증"

    path1 = convert_service._ensure_cli_recoded_tmp("job1", job, "기출.txt", text)
    assert path1.exists()
    assert path1.read_text(encoding="utf-8") == text
    assert job["_cli_recoded_path"] == str(path1)
    assert str(path1) in job["_extra_tmp_paths"]

    # 재호출 시 새로 만들지 않고 캐시된 경로를 반환한다(잡당 1회).
    path2 = convert_service._ensure_cli_recoded_tmp("job1", job, "기출.txt", text)
    assert path2 == path1
    assert len(job["_extra_tmp_paths"]) == 1


def test_ensure_cli_recoded_tmp_only_called_from_cli_branch_in_source():
    """codex·api 엔진은 `detected.text`를 직접 쓰므로 `_ensure_cli_recoded_tmp` 호출
    경로 자체가 없어야 한다 — `_do_convert` 전체에서 호출 지점이 정확히 1곳(CLI 분기)
    뿐임을 소스 검색으로 고정한다(회귀 방지, 검토 9)."""
    import inspect

    src = inspect.getsource(convert_service._do_convert)
    assert src.count("_ensure_cli_recoded_tmp(") == 1


# ---------------------------------------------------------------------------
# S16 검토 10 — `_do_fetch`의 file_mode(qnet 대표 파일)도 동일 판별 게이트를 통과한다.
# ---------------------------------------------------------------------------
class _FakeFileModeAdapter(Adapter):
    id = "fake"
    name = "fake"

    def __init__(self, fetched: FetchedFile) -> None:
        self._fetched = fetched

    def is_available(self, client):
        return True, None

    def fetch_exam(self, cert_ref, exam_ref, client, on_activity=None):
        return self._fetched


def test_do_fetch_file_mode_routes_by_detected_group_not_extension(monkeypatch, isolated_sources):
    """검토 10 — file_mode 대표 파일도 판별 계층을 통과해 group 기반으로 콘텐츠 블록을
    만든다. 확장자가 `.pdf`가 아닌데 매직 바이트는 PDF인 파일(실사용은 아니지만 게이트
    구조 완성 확인)도 document 블록으로 가야 한다(확장자 재분기 없음)."""
    from services.fetchers import registry as fetch_registry

    fetched = FetchedFile(
        filename="representative.download",  # 확장자가 .pdf가 아님
        data=b"%PDF-1.4 fake pdf body",
        content_type="application/octet-stream",
        cert_name="테스트자격증",
        exam_key="2024-1",
        exam_label="2024년 1회",
        level_hint="필기",
    )
    fake_adapter = _FakeFileModeAdapter(fetched)
    monkeypatch.setattr(fetch_registry, "get_adapter", lambda aid: fake_adapter)
    monkeypatch.setattr(fetch_registry, "new_client", lambda on_activity=None: object())

    captured: dict = {}

    def _fake_run_api_streaming(prompt_text, *, file_blocks, timeout_seconds, job_id, model):
        captured["file_blocks"] = file_blocks
        return '{"format_version": 1, "documents": []}'

    class _FakePreview:
        preview_id = "imp_fake"

    monkeypatch.setattr(convert_service, "_run_api_streaming", _fake_run_api_streaming)
    monkeypatch.setattr(
        convert_service.import_service, "create_preview", lambda *a, **k: _FakePreview()
    )

    job = convert_service._new_job_base(
        "fetch",
        resolved_engine=convert_service.llm_engine_service.ENGINE_CLAUDE_API,
        requested_engine="auto",
        api_model=convert_service.llm_engine_service.DEFAULT_API_MODEL,
    )
    job.update(
        {
            "_timeout": 60,
            "_adapter_id": "fake",
            "_cert_ref": "cert",
            "_exam_ref": "exam",
            "_fetch_source_url": None,
            "_exam_key_override": None,
        }
    )

    result = convert_service._do_fetch("ftc_test", job)

    assert result == {"result_preview_id": "imp_fake"}
    assert captured["file_blocks"][0]["type"] == "document"  # 확장자 재분기였다면 텍스트 블록
    saved = list(isolated_sources.glob("*representative.download"))
    assert saved == []  # pdf 판별 통과분 — C군 거부 저장 경로를 타지 않는다
