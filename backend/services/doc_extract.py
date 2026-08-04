"""B군(docx·xlsx) 서버측 텍스트 추출 (S16 — F42, 설계 §4.18 ④).

파서 채택(포맷당 1개, D1 확정): docx = **python-docx**, xlsx = **openpyxl**. 이 모듈은
바이트 → 직렬화 텍스트만 다루는 순수 추출기다 — convert 파이프라인(`convert_service`)과
원문 대조(`source_match.extract_source_text`, §4.17 ⑥)가 **같은 함수를 재사용**한다(분리만
한다 — 대조 로직 변경·추출 텍스트 저장·색인은 없음, 과설계 금지).

파서 예외는 두 typed 예외로만 좁혀 노출한다(원문 예외 텍스트는 로그에만 남긴다):
  - `DocExtractError` — 구조 해석 실패(암호·손상·구버전 포맷 위장). 판별이 docx/xlsx로
    확정된 뒤의 추출 실패 한정 — 호출부(`convert_service`)가 `error_info.kind:'parse_failed'`
    로 매핑한다(신규 kind 없음, §4.18 ④ 경계 규칙).
  - `DocTooLargeError` — xlsx 시트10/행500/열50 상한 초과 또는 추출 텍스트 200,000자 상한
    초과. 호출부가 `error_info.kind:'too_large'`로 매핑한다(§4.18 ⑤).
"""
from __future__ import annotations

import io
import logging
from typing import Any, List, Optional

_LOGGER = logging.getLogger(__name__)

# --- 상한(설계 §4.18 ④·⑤ 확정 — 임의 변경 금지) -----------------------------------
MAX_TEXT_CHARS = 200_000  # D2-⑥ — A군 디코드·B군 추출 텍스트 공통 상한
XLSX_MAX_SHEETS = 10
XLSX_MAX_ROWS = 500
XLSX_MAX_COLS = 50

_SPLIT_ACTION = "원본을 과목·회차 단위로 나눠 개별 파일로 반입해 주세요."


class DocExtractError(Exception):
    """구조 해석 실패(암호·손상·구버전 포맷 위장) — 판별이 docx/xlsx로 확정된 뒤의 추출
    실패 한정. `parse_failed`로 매핑된다(§4.18 ④ 경계 규칙)."""

    def __init__(self, message: str) -> None:
        super().__init__(message)
        self.public_message = message


class DocTooLargeError(Exception):
    """xlsx 구조 상한(시트10·행500·열50) 또는 추출 텍스트 200,000자 상한 초과 —
    `too_large`로 매핑된다(§4.18 ⑤. 부분 잘림 투입 없이 실행 전 중단)."""

    def __init__(self, message: str, *, action: str = _SPLIT_ACTION) -> None:
        super().__init__(message)
        self.public_message = message
        self.action = action


def enforce_max_chars(
    text: str, *, label: str = "추출된 텍스트", max_chars: Optional[int] = None
) -> None:
    """A군 디코드·B군 추출 텍스트 공통 200,000자 상한(§4.18 ⑤) — LLM 호출 전 비용 0 차단.
    서버측 자동 분할은 하지 않는다(F40-④ 원칙과 동일 — 분할은 사용자 몫).

    `max_chars`(S23 — F49, 설계 §4.25) — 기본 `None`이면 `MAX_TEXT_CHARS`(기존 동작 완전
    불변). 분할 반입(`split_service`)만 이 상한을 우회하려고 훨씬 큰 값을 넘겨 판별·추출
    계층 자체는 중복 구현하지 않고 재사용한다 — 값 자체(200,000)를 바꾸는 것이 아니라
    **호출자가 다른 상한을 적용**하는 것뿐이다."""
    limit = MAX_TEXT_CHARS if max_chars is None else max_chars
    if len(text) > limit:
        raise DocTooLargeError(
            f"{label}가 너무 깁니다(약 {len(text):,}자 — 상한 {limit:,}자)"
        )


# ---------------------------------------------------------------------------
# Markdown 표 직렬화 유틸 (docx 표·xlsx 시트 공용)
# ---------------------------------------------------------------------------
def _escape_md_cell(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\\", "\\\\").replace("|", "\\|")
    text = text.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    return text.strip()


def _md_table(rows: List[List[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# docx 추출 (python-docx)
# ---------------------------------------------------------------------------
def _iter_docx_blocks(document: Any):
    """문단·표를 **문서 순서 그대로** 순회 — python-docx가 `document.paragraphs`/
    `document.tables`를 따로 노출해 원 순서 정보가 없으므로 본문 XML을 직접 순회한다
    (python-docx 공식 문서가 권하는 관용구)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _docx_table_to_md(table: Any) -> str:
    rows = [[_escape_md_cell(cell.text) for cell in row.cells] for row in table.rows]
    rows = [r for r in rows if any(r)]
    return _md_table(rows)


def extract_docx_text(data: bytes, *, max_chars: Optional[int] = None) -> str:
    """docx 바이트 → 문단+표(Markdown 표) 텍스트를 **문서 순서대로** 직렬화(§4.18 ④).

    수식(OMML)·이미지·텍스트박스는 소실된다 — 감지는 `docx_has_lost_elements`가 별도로
    담당한다(호출부가 잡 notes에 소표기 1건을 붙인다). 200,000자 상한을 여기서 확인한다.
    `max_chars`(S23 — F49)는 `enforce_max_chars`로 그대로 전달(기본 동작 불변)."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - requirements에 고정 설치되어야 정상
        raise DocExtractError("서버에 문서 처리 라이브러리가 설치되어 있지 않습니다.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - 손상·암호·구조 위장 전부 흡수
        _LOGGER.info("docx 파싱 실패: %s", exc)
        raise DocExtractError(
            "워드(docx) 문서를 열 수 없습니다(손상되었거나 암호가 걸려 있을 수 있습니다)."
        ) from exc

    parts: List[str] = []
    try:
        for block in _iter_docx_blocks(document):
            if hasattr(block, "text"):  # Paragraph
                if block.text.strip():
                    parts.append(block.text)
            else:  # Table
                md = _docx_table_to_md(block)
                if md:
                    parts.append(md)
    except DocExtractError:
        raise
    except Exception as exc:  # noqa: BLE001 - 본문 순회 중 예외도 구조화한다
        _LOGGER.info("docx 본문 추출 실패: %s", exc)
        raise DocExtractError("워드(docx) 문서 내용을 추출하지 못했습니다.") from exc

    text = "\n\n".join(parts).strip()
    enforce_max_chars(text, max_chars=max_chars)  # S16 검토 6 — §4.18 ⑤ 확정 문안 그대로
    return text


def docx_has_lost_elements(data: bytes) -> bool:
    """수식(OMML)·그림·텍스트박스 등 텍스트 추출에서 소실되는 요소 감지(§4.18 ④ notes 소표기).

    docx는 zip이므로 `word/document.xml` 원문에서 태그 존재만 가볍게 검사한다(휴리스틱 —
    오탐보다 누락 방지가 목적). 감지 자체가 실패해도 오류로 취급하지 않고 False로 흡수한다
    (소표기 생략일 뿐 반입을 막지 않는다)."""
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    except Exception:  # noqa: BLE001
        return False
    markers = ("<m:oMath", "<w:drawing", "<w:pict", "<w:object")
    return any(marker in xml for marker in markers)


# ---------------------------------------------------------------------------
# xlsx 추출 (openpyxl)
# ---------------------------------------------------------------------------
def extract_xlsx_text(data: bytes, *, max_chars: Optional[int] = None) -> str:
    """xlsx 바이트 → 시트별 Markdown 표(`## 시트: {이름}` + 표, D2-③ 확정)로 직렬화.

    시트 10개·시트당 500행·행당 50열 상한 초과 시 **부분 잘림 없이** `DocTooLargeError`로
    중단한다(초과 지점을 message에 명시 — 이 구조 상한은 `max_chars`와 무관하게 항상
    적용된다, S23 F49 범위 밖 "B군 구조 상한 초과 분할 없음"). 수식 셀은 계산값
    (`data_only=True` — 캐시 없으면 빈 값). 200,000자 상한(텍스트 길이)만 `max_chars`로
    우회 가능하다(기본 `None` = 기존 동작 불변)."""
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover - requirements에 고정 설치되어야 정상
        raise DocExtractError("서버에 문서 처리 라이브러리가 설치되어 있지 않습니다.") from exc

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    except Exception as exc:  # noqa: BLE001 - 손상·암호·구조 위장 전부 흡수
        _LOGGER.info("xlsx 파싱 실패: %s", exc)
        raise DocExtractError(
            "엑셀(xlsx) 문서를 열 수 없습니다(손상되었거나 암호가 걸려 있을 수 있습니다)."
        ) from exc

    try:
        sheet_names = workbook.sheetnames
        if len(sheet_names) > XLSX_MAX_SHEETS:
            raise DocTooLargeError(
                f"시트 수({len(sheet_names)}개)가 시트 상한({XLSX_MAX_SHEETS}개)을 초과했습니다"
            )

        parts: List[str] = []
        for name in sheet_names:
            ws = workbook[name]
            # S16 검토 4 — `ws.max_row`/`max_column`은 서식만 있는 빈 셀도 "사용된 범위"로
            # 잡아 상한 오탐을 낸다(실측: 데이터 3행인데 600행 서식만으로 max_row=600).
            # **직렬화 루프가 실제로 세는 값**(값이 있는 행·그 안의 최대 열 인덱스)으로만
            # 판정한다 — 판정 기준과 직렬화 기준을 일치시킨다.
            raw_rows = [
                row for row in ws.iter_rows(values_only=True) if any(v is not None for v in row)
            ]
            if len(raw_rows) > XLSX_MAX_ROWS:
                raise DocTooLargeError(
                    f"시트 '{name}'의 {len(raw_rows)}행에서 행 상한({XLSX_MAX_ROWS})을 초과했습니다"
                )
            last_col = 0
            for row in raw_rows:
                for idx in range(len(row) - 1, -1, -1):
                    if row[idx] is not None:
                        last_col = max(last_col, idx + 1)
                        break
            if last_col > XLSX_MAX_COLS:
                raise DocTooLargeError(
                    f"시트 '{name}'의 열이 {last_col}개로 열 상한({XLSX_MAX_COLS})을 초과했습니다"
                )
            parts.append(f"## 시트: {name}")
            rows_out: List[List[str]] = [
                [_escape_md_cell(v) for v in row[:last_col]] for row in raw_rows
            ]
            if rows_out:
                parts.append(_md_table(rows_out))
    except DocExtractError:
        raise
    except DocTooLargeError:
        raise
    except Exception as exc:  # noqa: BLE001 - 시트 순회 중 예외도 구조화한다
        _LOGGER.info("xlsx 시트 추출 실패: %s", exc)
        raise DocExtractError("엑셀(xlsx) 문서 내용을 추출하지 못했습니다.") from exc
    finally:
        workbook.close()

    text = "\n\n".join(parts).strip()
    enforce_max_chars(text, max_chars=max_chars)  # S16 검토 6 — §4.18 ⑤ 확정 문안 그대로
    return text
